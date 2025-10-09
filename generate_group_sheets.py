"""
Script pour générer automatiquement les feuilles de présence, de contenu et de suivi des paiements
pour un groupe Supabase.

Génère 3 documents par groupe :
1. Feuille de présence avec 12 séances
2. Feuille de contenu avec 6 séances (2 pages)
3. Feuille de suivi des paiements

Usage:
    python generate_group_sheets.py <group_id>

Exemple:
    python generate_group_sheets.py 1
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from supabase import create_client
from datetime import datetime
from dotenv import load_dotenv

# Charger les variables d'environnement
load_dotenv()


def get_supabase_client():
    """Crée et retourne un client Supabase."""
    url = os.getenv("SUPABASE_URL")
    key = os.getenv("SUPABASE_KEY")

    if not url or not key:
        raise ValueError("SUPABASE_URL et SUPABASE_KEY doivent être définis dans .env ou secrets.toml")

    return create_client(url, key)


def set_cell_border(cell, **kwargs):
    """
    Définit les bordures d'une cellule de tableau.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Liste des côtés de bordure
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '12')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)


def get_group_data(group_id):
    """
    Récupère toutes les données d'un groupe.
    """
    supabase = get_supabase_client()

    print(f"Récupération des données pour le groupe {group_id}...")

    # Récupérer les informations du groupe avec la langue
    group_response = supabase.table('groups').select(
        '*, languages(name)'
    ).eq('id', group_id).execute()

    if not group_response.data:
        print(f"❌ Groupe {group_id} non trouvé")
        return None

    group_data = group_response.data[0]
    print(f"✓ Groupe trouvé: {group_data['name']}")

    # Récupérer le planning du groupe avec les salles
    schedule_response = supabase.table('schedule').select(
        '*, classrooms(name, location)'
    ).eq('group_id', group_id).execute()

    # Récupérer les enseignants du groupe
    teachers_response = supabase.table('group_teacher').select(
        '*, teachers(first_name, last_name)'
    ).eq('group_id', group_id).execute()

    teachers_names = []
    if teachers_response.data:
        for gt in teachers_response.data:
            teacher = gt.get('teachers', {})
            if teacher:
                teachers_names.append(f"{teacher['first_name']} {teacher['last_name']}")

    # Récupérer les étudiants actifs avec informations de paiement
    enrollments_response = supabase.table('enrollments').select(
        '*, students(id, first_name, last_name, student_code)'
    ).eq('group_id', group_id).eq('enrollment_active', True).execute()

    students_list = []
    if enrollments_response.data:
        for enr in enrollments_response.data:
            student = enr.get('students', {})
            if student:
                student_id = student['id']

                # Récupérer les paiements de l'étudiant (triés par date)
                payments_response = supabase.table('payments').select('*').eq(
                    'student_id', student_id
                ).order('payment_date').execute()

                # Liste des paiements individuels
                payment_list = []
                if payments_response.data:
                    for payment in payments_response.data[:3]:  # Limiter à 3 premiers paiements
                        payment_list.append(payment['amount'])

                total_paid = sum([p['amount'] for p in payments_response.data]) if payments_response.data else 0
                total_course_fee = enr.get('total_course_fee', 0)
                remaining = total_course_fee - total_paid

                # Frais d'inscription (on considère qu'il faut au moins 1000 DA)
                registration_paid = total_paid >= 1000

                # Si tout payé en une fois
                paid_in_full = total_paid >= total_course_fee and len(payments_response.data) == 1

                students_list.append({
                    'first_name': student['first_name'],
                    'last_name': student['last_name'],
                    'student_code': student.get('student_code', 'N/A'),
                    'total_course_fee': total_course_fee,
                    'total_paid': total_paid,
                    'remaining': remaining,
                    'registration_paid': registration_paid,
                    'paid_in_full': paid_in_full,
                    'payments': payment_list  # Liste des paiements individuels
                })

    print(f"✓ {len(students_list)} étudiant(s) actif(s)")

    # Formater les informations de planning
    schedule_info = []
    if schedule_response.data:
        day_translation = {
            'Mon': 'Lundi',
            'Tue': 'Mardi',
            'Wed': 'Mercredi',
            'Thu': 'Jeudi',
            'Fri': 'Vendredi',
            'Sat': 'Samedi',
            'Sun': 'Dimanche'
        }

        for sched in schedule_response.data:
            day_fr = day_translation.get(sched['day_of_week'], sched['day_of_week'])
            classroom_name = 'En ligne' if sched.get('is_online') else (
                sched.get('classrooms', {}).get('name', 'N/A') if sched.get('classrooms') else 'N/A'
            )

            schedule_info.append({
                'day': day_fr,
                'start_time': sched['start_time'],
                'end_time': sched['end_time'],
                'classroom': classroom_name
            })

    # Déterminer le mode de cours
    mode = group_data.get('mode', '')
    is_online = 'online' in mode.lower()

    return {
        'group_name': group_data.get('name', 'N/A'),
        'language': group_data.get('languages', {}).get('name', 'N/A'),
        'level': group_data.get('level', 'N/A'),
        'teachers': ', '.join(teachers_names) if teachers_names else 'N/A',
        'start_date': group_data.get('start_date', 'N/A'),
        'schedule': schedule_info,
        'students': students_list,
        'is_online': is_online
    }


def create_attendance_sheet(group_data):
    """
    Crée une feuille de présence en mode paysage avec 12 séances.
    """
    doc = Document()

    # Configuration de la page en paysage
    section = doc.sections[0]
    section.page_height = Inches(8.27)  # A4 largeur devient hauteur
    section.page_width = Inches(11.69)  # A4 hauteur devient largeur
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('FEUILLE DE PRÉSENCE 2025/2026')
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph()  # Espace

    # En-tête avec informations du groupe
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Light Grid Accent 1'

    # Ligne 1
    info_table.rows[0].cells[0].text = 'Enseignant:'
    info_table.rows[0].cells[1].text = group_data['teachers']
    info_table.rows[0].cells[2].text = 'Groupe:'
    info_table.rows[0].cells[3].text = group_data['group_name']

    # Ligne 2
    info_table.rows[1].cells[0].text = 'Langue:'
    info_table.rows[1].cells[1].text = group_data['language']
    info_table.rows[1].cells[2].text = 'Niveau:'
    info_table.rows[1].cells[3].text = str(group_data['level'])

    # Ligne 3
    if group_data['schedule']:
        sched = group_data['schedule'][0]
        info_table.rows[2].cells[0].text = 'Jour:'
        info_table.rows[2].cells[1].text = f"{sched['day']} {sched['start_time']}-{sched['end_time']}"
        info_table.rows[2].cells[2].text = 'Salle:'
        info_table.rows[2].cells[3].text = sched['classroom']

    # Appliquer la taille de police 11pt à toutes les cellules d'info
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()  # Espace

    # Date de début
    if group_data['start_date'] != 'N/A':
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(f"Date de début: {group_data['start_date']}")
        date_run.font.size = Pt(11)
        date_run.font.bold = True

    doc.add_paragraph()  # Espace

    # Tableau de présence
    num_students = len(group_data['students'])
    attendance_table = doc.add_table(rows=num_students + 2, cols=14)
    attendance_table.style = 'Light Grid Accent 1'

    # En-tête du tableau
    hdr_cells = attendance_table.rows[0].cells
    hdr_cells[0].text = 'Nom et Prénom'
    hdr_cells[0].merge(hdr_cells[1])

    for i in range(12):
        cell = attendance_table.rows[0].cells[i + 2]
        cell.text = f'S{i + 1}'
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Ligne "Date"
    date_row = attendance_table.rows[1]
    date_row.cells[0].text = 'Date'
    date_row.cells[0].merge(date_row.cells[1])

    for i in range(12):
        cell = date_row.cells[i + 2]
        cell.text = ''  # À remplir manuellement
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remplir avec les étudiants
    for idx, student in enumerate(group_data['students']):
        row = attendance_table.rows[idx + 2]
        full_name = f"{student['last_name']} {student['first_name']}"
        row.cells[0].text = full_name
        row.cells[0].merge(row.cells[1])

        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

        # Cellules de présence vides
        for i in range(12):
            cell = row.cells[i + 2]
            cell.text = ''

    # Ajuster la largeur des colonnes
    # Colonne nom plus large
    for row in attendance_table.rows:
        row.cells[0].width = Cm(4)

    return doc


def create_content_sheet(group_data):
    """
    Crée une feuille de contenu des séances avec 6 séances (2 pages).
    """
    doc = Document()

    # Configuration de la page
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('FEUILLE DE CONTENU DES SÉANCES 2025/2026')
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph()  # Espace

    # En-tête avec informations du groupe
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Light Grid Accent 1'

    # Ligne 1
    info_table.rows[0].cells[0].text = 'Enseignant:'
    info_table.rows[0].cells[1].text = group_data['teachers']
    info_table.rows[0].cells[2].text = 'Groupe:'
    info_table.rows[0].cells[3].text = group_data['group_name']

    # Ligne 2
    info_table.rows[1].cells[0].text = 'Langue:'
    info_table.rows[1].cells[1].text = group_data['language']
    info_table.rows[1].cells[2].text = 'Niveau:'
    info_table.rows[1].cells[3].text = str(group_data['level'])

    # Ligne 3
    if group_data['schedule']:
        sched = group_data['schedule'][0]
        info_table.rows[2].cells[0].text = 'Jour:'
        info_table.rows[2].cells[1].text = f"{sched['day']} {sched['start_time']}-{sched['end_time']}"
        info_table.rows[2].cells[2].text = 'Salle:'
        info_table.rows[2].cells[3].text = sched['classroom']

    # Appliquer la taille de police 11pt
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()  # Espace

    # Date de début
    if group_data['start_date'] != 'N/A':
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(f"Date de début: {group_data['start_date']}")
        date_run.font.size = Pt(11)
        date_run.font.bold = True

    doc.add_paragraph()  # Espace

    # Créer 6 séances (3 par page)
    for i in range(6):
        # Titre de la séance
        session_title = doc.add_paragraph()
        session_title_run = session_title.add_run(f'Séance {i + 1}')
        session_title_run.font.size = Pt(13)
        session_title_run.font.bold = True

        # Date de la séance
        date_para = doc.add_paragraph()
        date_run = date_para.add_run('Date: _________________')
        date_run.font.size = Pt(11)

        # Grande zone de contenu avec lignes
        content_table = doc.add_table(rows=1, cols=1)
        content_table.style = 'Light Grid Accent 1'

        cell = content_table.rows[0].cells[0]

        # Ajouter plusieurs lignes vides pour le contenu
        for j in range(8):
            para = cell.add_paragraph('_' * 100)
            for run in para.runs:
                run.font.size = Pt(10)

        # Définir une hauteur minimale pour la cellule
        cell.height = Cm(5)

        # Espace entre les séances
        if i < 5:  # Pas d'espace après la dernière séance
            doc.add_paragraph()

        # Saut de page après 3 séances
        if i == 2:
            doc.add_page_break()

    return doc


def create_payment_tracking_sheet(group_data):
    """
    Crée une feuille de suivi des paiements en mode paysage.
    """
    doc = Document()

    # Configuration de la page en paysage
    section = doc.sections[0]
    section.page_height = Inches(8.27)  # A4 largeur devient hauteur
    section.page_width = Inches(11.69)  # A4 hauteur devient largeur
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('SUIVI DES PAIEMENTS 2025/2026')
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph()  # Espace

    # En-tête avec informations du groupe
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Light Grid Accent 1'

    # Ligne 1
    info_table.rows[0].cells[0].text = 'Enseignant:'
    info_table.rows[0].cells[1].text = group_data['teachers']
    info_table.rows[0].cells[2].text = 'Groupe:'
    info_table.rows[0].cells[3].text = group_data['group_name']

    # Ligne 2
    info_table.rows[1].cells[0].text = 'Langue:'
    info_table.rows[1].cells[1].text = group_data['language']
    info_table.rows[1].cells[2].text = 'Niveau:'
    info_table.rows[1].cells[3].text = str(group_data['level'])

    # Ligne 3
    if group_data['schedule']:
        sched = group_data['schedule'][0]
        info_table.rows[2].cells[0].text = 'Jour:'
        info_table.rows[2].cells[1].text = f"{sched['day']} {sched['start_time']}-{sched['end_time']}"
        info_table.rows[2].cells[2].text = 'Salle:'
        info_table.rows[2].cells[3].text = sched['classroom']

    # Appliquer la taille de police 11pt à toutes les cellules d'info
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()  # Espace

    # Date de début
    if group_data['start_date'] != 'N/A':
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(f"Date de début: {group_data['start_date']}")
        date_run.font.size = Pt(11)
        date_run.font.bold = True

    doc.add_paragraph()  # Espace

    # Tableau de suivi des paiements
    num_students = len(group_data['students'])
    # Colonnes: Nom | Code | Frais Total | Frais Insc. | M1 | M2 | M3 | Total Payé | Restant | Observations
    payment_table = doc.add_table(rows=num_students + 1, cols=10)
    payment_table.style = 'Light Grid Accent 1'

    # En-tête du tableau
    headers = ['Nom et Prénom', 'Code', 'Frais Total', 'Frais Insc.', 'M1', 'M2', 'M3', 'Total Payé', 'Restant', 'Observations']
    hdr_cells = payment_table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Remplir avec les étudiants
    for idx, student in enumerate(group_data['students']):
        row = payment_table.rows[idx + 1]

        # Nom et Prénom
        full_name = f"{student['last_name']} {student['first_name']}"
        row.cells[0].text = full_name

        # Code étudiant
        row.cells[1].text = student['student_code']

        # Frais total
        row.cells[2].text = f"{student['total_course_fee']:,.0f} DA"

        # Frais d'inscription
        if student['registration_paid']:
            row.cells[3].text = "✓ Payé"
        else:
            row.cells[3].text = "✗ Non payé"

        # Mensualités (M1, M2, M3)
        if student['paid_in_full']:
            # Si tout payé en une fois, barrer les 3 cellules de mensualités
            for m_idx in range(4, 7):  # Colonnes M1, M2, M3
                para = row.cells[m_idx].paragraphs[0]
                run = para.add_run("X")
                run.font.strike = True
                run.font.size = Pt(9)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Remplir avec les paiements réels
            payments = student.get('payments', [])
            for m_idx in range(4, 7):  # Colonnes M1, M2, M3
                payment_idx = m_idx - 4  # 0, 1, 2
                if payment_idx < len(payments):
                    # Afficher le montant payé
                    row.cells[m_idx].text = f"{payments[payment_idx]:,.0f}"
                    for paragraph in row.cells[m_idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                else:
                    # Laisser vide si pas encore de paiement
                    row.cells[m_idx].text = ""

        # Total payé
        row.cells[7].text = f"{student['total_paid']:,.0f} DA"

        # Restant
        if student['remaining'] <= 0:
            row.cells[8].text = "SOLDÉ"
            for paragraph in row.cells[8].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        else:
            row.cells[8].text = f"{student['remaining']:,.0f} DA"

        # Observations (vide, à remplir manuellement)
        row.cells[9].text = ""

        # Appliquer la taille de police
        for i in range(10):
            if i not in [4, 5, 6]:  # Sauf les colonnes mensualités déjà formatées
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Ajuster les largeurs des colonnes
    for row in payment_table.rows:
        row.cells[0].width = Cm(3.2)  # Nom
        row.cells[1].width = Cm(1.5)  # Code
        row.cells[2].width = Cm(2)    # Frais Total
        row.cells[3].width = Cm(1.8)  # Frais Insc.
        row.cells[4].width = Cm(1.5)  # M1
        row.cells[5].width = Cm(1.5)  # M2
        row.cells[6].width = Cm(1.5)  # M3
        row.cells[7].width = Cm(2)    # Total Payé
        row.cells[8].width = Cm(1.8)  # Restant
        row.cells[9].width = Cm(3.5)  # Observations

    return doc


def generate_group_sheets(group_id, output_folder='feuilles_groupe'):
    """
    Génère les feuilles de présence, de contenu et de suivi des paiements pour un groupe.

    Args:
        group_id: ID du groupe
        output_folder: Dossier de sortie pour les feuilles générées
    """
    print(f"\n{'='*60}")
    print(f"Génération des feuilles pour le groupe {group_id}")
    print(f"{'='*60}\n")

    # Récupérer les données
    group_data = get_group_data(group_id)

    if group_data is None:
        return

    if not group_data['students']:
        print("⚠ Aucun étudiant actif dans ce groupe")
        print("  Les feuilles seront générées sans étudiants.\n")

    # Créer le dossier de sortie
    os.makedirs(output_folder, exist_ok=True)
    print(f"📁 Dossier de sortie: {os.path.abspath(output_folder)}\n")

    print("Génération des documents...\n")

    # Générer la feuille de présence
    print("  [1/3] Création de la feuille de présence...")
    attendance_doc = create_attendance_sheet(group_data)
    attendance_filename = f"Presence_{group_data['group_name'].replace(' ', '_')}.docx"
    attendance_filepath = os.path.join(output_folder, attendance_filename)
    attendance_doc.save(attendance_filepath)
    print(f"        ✓ {attendance_filename}")

    # Générer la feuille de contenu
    print("  [2/3] Création de la feuille de contenu...")
    content_doc = create_content_sheet(group_data)
    content_filename = f"Contenu_{group_data['group_name'].replace(' ', '_')}.docx"
    content_filepath = os.path.join(output_folder, content_filename)
    content_doc.save(content_filepath)
    print(f"        ✓ {content_filename}")

    # Générer la feuille de suivi des paiements
    print("  [3/3] Création de la feuille de suivi des paiements...")
    payment_doc = create_payment_tracking_sheet(group_data)
    payment_filename = f"Paiements_{group_data['group_name'].replace(' ', '_')}.docx"
    payment_filepath = os.path.join(output_folder, payment_filename)
    payment_doc.save(payment_filepath)
    print(f"        ✓ {payment_filename}")

    print(f"\n{'='*60}")
    print(f"✅ 3 document(s) générés avec succès!")
    print(f"📂 Emplacement: {os.path.abspath(output_folder)}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_group_sheets.py <group_id> [output_folder]")
        print("\nExemples:")
        print("  python generate_group_sheets.py 1")
        print("  python generate_group_sheets.py 1 mes_feuilles")
        sys.exit(1)

    group_id = int(sys.argv[1])
    output_folder = sys.argv[2] if len(sys.argv) > 2 else 'feuilles_groupe'

    try:
        generate_group_sheets(group_id, output_folder)
    except Exception as e:
        print(f"\n❌ Erreur: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
