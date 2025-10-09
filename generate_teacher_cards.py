"""
Script pour générer automatiquement des fiches enseignant Word
pour tous les enseignants d'un groupe Supabase.

Usage:
    python generate_teacher_cards.py <group_id>

Exemple:
    python generate_teacher_cards.py 1
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from supabase import create_client
from datetime import datetime
from dotenv import load_dotenv

# Charger les variables d'environnement
load_dotenv()


def get_supabase_client():
    """Crée et retourne un client Supabase."""
    url = os.getenv("SUPABASE_URL")
    key = os.getenv("SUPABASE_KEY")

    if not url or not key:
        raise ValueError("SUPABASE_URL et SUPABASE_KEY doivent être définis dans .env ou secrets.toml")

    return create_client(url, key)


def get_group_teacher_data(group_id):
    """
    Récupère toutes les données nécessaires pour générer les fiches enseignants
    d'un groupe donné.
    """
    supabase = get_supabase_client()

    print(f"Récupération des données pour le groupe {group_id}...")

    # Récupérer les informations du groupe avec la langue
    group_response = supabase.table('groups').select(
        '*, languages(name)'
    ).eq('id', group_id).execute()

    if not group_response.data:
        print(f"❌ Groupe {group_id} non trouvé")
        return None

    group_data = group_response.data[0]
    print(f"✓ Groupe trouvé: {group_data['name']}")

    # Récupérer le planning du groupe
    schedule_response = supabase.table('schedule').select(
        '*'
    ).eq('group_id', group_id).execute()

    # Récupérer les enseignants du groupe
    teachers_response = supabase.table('group_teacher').select(
        '*, teachers(first_name, last_name, email)'
    ).eq('group_id', group_id).execute()

    if not teachers_response.data:
        print("❌ Aucun enseignant associé à ce groupe")
        return []

    print(f"✓ {len(teachers_response.data)} enseignant(s) trouvé(s)")

    # Déterminer le mode de cours
    mode = group_data.get('mode', '')
    mode_translation = {
        'online_group': 'En ligne - Groupe',
        'online_individual': 'En ligne - Individuel',
        'presential_group': 'Présentiel - Groupe',
        'presential_individual': 'Présentiel - Individuel'
    }
    course_mode = mode_translation.get(mode, mode)

    # Récupérer les informations de planning
    schedule_info = []
    if schedule_response.data:
        day_translation = {
            'Mon': 'Lundi',
            'Tue': 'Mardi',
            'Wed': 'Mercredi',
            'Thu': 'Jeudi',
            'Fri': 'Vendredi',
            'Sat': 'Samedi',
            'Sun': 'Dimanche'
        }

        for sched in schedule_response.data:
            day_fr = day_translation.get(sched['day_of_week'], sched['day_of_week'])
            schedule_info.append({
                'day': day_fr,
                'start_time': sched['start_time'],
                'end_time': sched['end_time']
            })

    teachers_data = []

    for gt in teachers_response.data:
        teacher = gt.get('teachers', {})
        if teacher:
            teacher_data = {
                'first_name': teacher['first_name'],
                'last_name': teacher['last_name'],
                'full_name': f"{teacher['first_name']} {teacher['last_name']}",
                'email': teacher.get('email', 'N/A'),
                'group_name': group_data.get('name', 'N/A'),
                'language': group_data.get('languages', {}).get('name', 'N/A'),
                'course_mode': course_mode,
                'start_date': group_data.get('start_date', None),
                'schedule': schedule_info
            }
            teachers_data.append(teacher_data)
            print(f"  ✓ {teacher_data['full_name']}")

    return teachers_data


def create_teacher_card(teacher_data, header_image_path='header_image.png'):
    """
    Crée un document Word de fiche enseignant.

    Args:
        teacher_data: Dictionnaire contenant les données de l'enseignant
        header_image_path: Chemin vers l'image d'en-tête (par défaut: 'header_image.png')
    """
    doc = Document()

    # Configuration de la page
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Ajouter l'image d'en-tête si elle existe
    if os.path.exists(header_image_path):
        # Calculer la largeur utilisable (largeur de page - marges)
        usable_width = section.page_width - section.left_margin - section.right_margin

        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.add_run()
        run.add_picture(header_image_path, width=usable_width)

        doc.add_paragraph()  # Espace après l'image

    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace

    # Année scolaire (centré)
    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_run = year_para.add_run("Année Scolaire 2025/2026")
    year_run.font.size = Pt(18)
    year_run.font.bold = True

    doc.add_paragraph()  # Espace

    # Nom de l'enseignant (très grand, centré)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(teacher_data['full_name'])
    name_run.font.size = Pt(36)
    name_run.font.bold = True

    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace

    # Groupe (grand, centré)
    group_para = doc.add_paragraph()
    group_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    group_run = group_para.add_run(f"Groupe: {teacher_data['group_name']}")
    group_run.font.size = Pt(28)

    doc.add_paragraph()  # Espace

    # Mode de cours (grand, centré)
    mode_para = doc.add_paragraph()
    mode_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mode_run = mode_para.add_run(teacher_data['course_mode'])
    mode_run.font.size = Pt(24)

    doc.add_paragraph()  # Espace

    # Date de début (si disponible)
    if teacher_data.get('start_date'):
        start_date_para = doc.add_paragraph()
        start_date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        start_date_run = start_date_para.add_run(f"Début: {teacher_data['start_date']}")
        start_date_run.font.size = Pt(22)

    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace

    # Planning (grand, centré)
    if teacher_data['schedule']:
        for sched in teacher_data['schedule']:
            # Jour
            day_para = doc.add_paragraph()
            day_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            day_run = day_para.add_run(sched['day'])
            day_run.font.size = Pt(26)
            day_run.font.bold = True

            # Horaires
            time_para = doc.add_paragraph()
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            time_run = time_para.add_run(f"{sched['start_time']} - {sched['end_time']}")
            time_run.font.size = Pt(24)

            doc.add_paragraph()  # Espace entre les sessions
    else:
        no_schedule_para = doc.add_paragraph()
        no_schedule_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        no_schedule_run = no_schedule_para.add_run("Planning non défini")
        no_schedule_run.font.size = Pt(24)

    return doc


def generate_teacher_cards(group_id, output_folder='fiches_enseignants', header_image='header_image.png'):
    """
    Génère les fiches enseignant pour tous les enseignants d'un groupe.

    Args:
        group_id: ID du groupe
        output_folder: Dossier de sortie pour les fiches générées
        header_image: Chemin vers l'image d'en-tête (par défaut: 'header_image.png')
    """
    print(f"\n{'='*60}")
    print(f"Génération des fiches enseignant pour le groupe {group_id}")
    print(f"{'='*60}\n")

    # Récupérer les données
    teachers_data = get_group_teacher_data(group_id)

    if teachers_data is None:
        return

    if not teachers_data:
        print("Aucune fiche à générer.")
        return

    # Créer le dossier de sortie
    os.makedirs(output_folder, exist_ok=True)
    print(f"\n📁 Dossier de sortie: {os.path.abspath(output_folder)}\n")

    # Vérifier si l'image d'en-tête existe
    if os.path.exists(header_image):
        print(f"✓ Image d'en-tête trouvée: {header_image}")
    else:
        print(f"⚠ Image d'en-tête non trouvée: {header_image}")
        print(f"  Les fiches seront générées sans en-tête image.\n")

    # Générer les fiches
    print("\nGénération des fiches Word...\n")

    for i, teacher_data in enumerate(teachers_data, 1):
        # Créer le document
        doc = create_teacher_card(teacher_data, header_image)

        # Nom du fichier
        filename = f"Fiche_Enseignant_{teacher_data['last_name']}_{teacher_data['first_name']}_{teacher_data['group_name'].replace(' ', '_')}.docx"
        filepath = os.path.join(output_folder, filename)

        # Sauvegarder
        doc.save(filepath)
        print(f"  [{i}/{len(teachers_data)}] ✓ {filename}")

    print(f"\n{'='*60}")
    print(f"✅ {len(teachers_data)} fiche(s) générée(s) avec succès!")
    print(f"📂 Emplacement: {os.path.abspath(output_folder)}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_teacher_cards.py <group_id> [output_folder] [header_image]")
        print("\nExemples:")
        print("  python generate_teacher_cards.py 1")
        print("  python generate_teacher_cards.py 1 mes_fiches")
        print("  python generate_teacher_cards.py 1 mes_fiches mon_logo.png")
        sys.exit(1)

    group_id = int(sys.argv[1])
    output_folder = sys.argv[2] if len(sys.argv) > 2 else 'fiches_enseignants'
    header_image = sys.argv[3] if len(sys.argv) > 3 else 'header_image.png'

    try:
        generate_teacher_cards(group_id, output_folder, header_image)
    except Exception as e:
        print(f"\n❌ Erreur: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
