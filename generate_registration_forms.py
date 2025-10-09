"""
Script pour générer automatiquement les fiches d'inscription Word
pour tous les étudiants d'un groupe Supabase.

Usage:
    python generate_registration_forms.py <group_id>

Exemple:
    python generate_registration_forms.py 1
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from supabase import create_client
from datetime import datetime
from dotenv import load_dotenv

# Charger les variables d'environnement
load_dotenv()


def get_supabase_client():
    """Crée et retourne un client Supabase."""
    url = os.getenv("SUPABASE_URL")
    key = os.getenv("SUPABASE_KEY")

    if not url or not key:
        raise ValueError("SUPABASE_URL et SUPABASE_KEY doivent être définis dans .env ou secrets.toml")

    return create_client(url, key)


def get_group_students_data(group_id):
    """
    Récupère toutes les données nécessaires pour générer les fiches d'inscription
    des étudiants d'un groupe donné.
    """
    supabase = get_supabase_client()

    print(f"Récupération des données pour le groupe {group_id}...")

    # Récupérer les informations du groupe avec la langue
    group_response = supabase.table('groups').select(
        '*, languages(name)'
    ).eq('id', group_id).execute()

    if not group_response.data:
        print(f"❌ Groupe {group_id} non trouvé")
        return None

    group_data = group_response.data[0]
    print(f"✓ Groupe trouvé: {group_data['name']}")

    # Récupérer les étudiants inscrits dans ce groupe via la table enrollments
    enrollments_response = supabase.table('enrollments').select(
        'student_id, level, enrollment_active, enrollment_date'
    ).eq('group_id', group_id).execute()

    if not enrollments_response.data:
        print("❌ Aucun étudiant inscrit dans ce groupe")
        return []

    print(f"✓ {len(enrollments_response.data)} étudiant(s) trouvé(s)")

    students_data = []

    for enrollment in enrollments_response.data:
        student_id = enrollment['student_id']

        # Récupérer les informations de l'étudiant
        student_response = supabase.table('students').select('*').eq('id', student_id).execute()

        if not student_response.data:
            continue

        student = student_response.data[0]

        # Récupérer l'email et le téléphone
        email = student.get('email', 'N/A')
        phone_number = student.get('phone_number', 'N/A')

        # Récupérer les paiements de l'étudiant
        # Note: Les paiements ne sont pas liés à un groupe spécifique dans le schéma
        payments_response = supabase.table('payments').select('*').eq(
            'student_id', student_id
        ).execute()

        total_paid = sum([p['amount'] for p in payments_response.data]) if payments_response.data else 0
        has_paid_minimum = total_paid >= 1000

        # Récupérer le planning du groupe avec les salles
        schedule_response = supabase.table('schedule').select(
            '*, classrooms(name, location)'
        ).eq('group_id', group_id).execute()

        # Récupérer les enseignants du groupe
        teachers_response = supabase.table('group_teacher').select(
            '*, teachers(first_name, last_name)'
        ).eq('group_id', group_id).execute()

        teachers_names = []
        if teachers_response.data:
            for gt in teachers_response.data:
                teacher = gt.get('teachers', {})
                if teacher:
                    teachers_names.append(f"{teacher['first_name']} {teacher['last_name']}")

        # Formater les informations de planning
        schedule_info = []
        if schedule_response.data:
            day_translation = {
                'Mon': 'Lundi',
                'Tue': 'Mardi',
                'Wed': 'Mercredi',
                'Thu': 'Jeudi',
                'Fri': 'Vendredi',
                'Sat': 'Samedi',
                'Sun': 'Dimanche'
            }

            for sched in schedule_response.data:
                day_fr = day_translation.get(sched['day_of_week'], sched['day_of_week'])
                start_time = sched['start_time']
                end_time = sched['end_time']

                classroom_name = 'En ligne' if sched.get('is_online') else (
                    sched.get('classrooms', {}).get('name', 'N/A') if sched.get('classrooms') else 'N/A'
                )

                schedule_info.append({
                    'day': day_fr,
                    'start_time': start_time,
                    'end_time': end_time,
                    'classroom': classroom_name
                })

        # Déterminer le type de cours à partir du mode
        mode = group_data.get('mode', '')
        mode_translation = {
            'online_group': 'En ligne - Groupe',
            'online_individual': 'En ligne - Individuel',
            'presential_group': 'Présentiel - Groupe',
            'presential_individual': 'Présentiel - Individuel'
        }
        course_mode = mode_translation.get(mode, mode)

        # Compiler toutes les données pour cet étudiant
        student_data = {
            'student_code': student.get('student_code', 'N/A'),
            'creation_date': enrollment.get('enrollment_date', datetime.now().isoformat())[:10],
            'first_name': student['first_name'],
            'last_name': student['last_name'],
            'full_name': f"{student['first_name']} {student['last_name']}",
            'email': email,
            'phone_number': phone_number,
            'date_of_birth': student.get('birth_date', 'N/A'),
            'has_paid_minimum': 'Oui' if has_paid_minimum else 'Non',
            'total_paid': total_paid,
            'course_mode': course_mode,
            'language': group_data.get('languages', {}).get('name', 'N/A'),
            'level': str(enrollment.get('level', group_data.get('level', 'N/A'))),
            'group_name': group_data.get('name', 'N/A'),
            'teachers': ', '.join(teachers_names) if teachers_names else 'N/A',
            'schedule': schedule_info
        }

        students_data.append(student_data)
        print(f"  ✓ {student_data['full_name']} ({student_data['student_code']})")

    return students_data


def create_registration_form(student_data, header_image_path='header_image.png'):
    """
    Crée un document Word de fiche d'inscription pour un étudiant.

    Args:
        student_data: Dictionnaire contenant les données de l'étudiant
        header_image_path: Chemin vers l'image d'en-tête (par défaut: 'header_image.png')
    """
    doc = Document()

    # Configuration de la page
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Ajouter l'image d'en-tête si elle existe
    if os.path.exists(header_image_path):
        # Calculer la largeur utilisable (largeur de page - marges)
        usable_width = section.page_width - section.left_margin - section.right_margin

        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.add_run()
        run.add_picture(header_image_path, width=usable_width)

        doc.add_paragraph()  # Espace après l'image
    else:
        # Si l'image n'existe pas, afficher le titre et sous-titre par défaut
        title = doc.add_heading('FICHE D\'INSCRIPTION 2025/2026', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph('Institut Torii - École de Langues Asiatiques')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.italic = True

        doc.add_paragraph()  # Espace

    # Informations de l'étudiant
    doc.add_heading('INFORMATIONS PERSONNELLES', 2)

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'

    # Remplir le tableau
    cells_data = [
        ('Code Étudiant:', student_data['student_code']),
        ('Nom complet:', student_data['full_name']),
        ('Email:', student_data['email']),
        ('Téléphone:', student_data['phone_number']),
        ('Date de naissance:', student_data['date_of_birth']),
        ('Date d\'inscription:', student_data['creation_date']),
        ('Paiement initial (≥1000 DA):', student_data['has_paid_minimum']),
        ('Montant total payé:', f"{student_data['total_paid']:,.0f} DA")
    ]

    for i, (label, value) in enumerate(cells_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        # Définir la taille de police à 12pt pour les deux cellules
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.font.size = Pt(12)
            run.font.bold = True
        for run in table.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(12)

    doc.add_paragraph()  # Espace

    # Informations du cours
    doc.add_heading('INFORMATIONS DU COURS', 2)

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Light Grid Accent 1'

    cells_data2 = [
        ('Langue:', student_data['language']),
        ('Niveau:', student_data['level']),
        ('Mode de cours:', student_data['course_mode']),
        ('Groupe:', student_data['group_name'])
    ]

    for i, (label, value) in enumerate(cells_data2):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = str(value)
        # Définir la taille de police à 12pt pour les deux cellules
        for run in table2.rows[i].cells[0].paragraphs[0].runs:
            run.font.size = Pt(12)
            run.font.bold = True
        for run in table2.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(12)

    doc.add_paragraph()  # Espace

    # Planning
    doc.add_heading('EMPLOI DU TEMPS', 2)

    if student_data['schedule']:
        # Ajouter les enseignants
        teachers_para = doc.add_paragraph()
        run1 = teachers_para.add_run('Enseignant(s): ')
        run1.bold = True
        run1.font.size = Pt(12)
        run2 = teachers_para.add_run(student_data['teachers'])
        run2.font.size = Pt(12)

        doc.add_paragraph()  # Espace

        # Tableau du planning
        schedule_table = doc.add_table(rows=len(student_data['schedule']) + 1, cols=4)
        schedule_table.style = 'Light Grid Accent 1'

        # En-têtes
        headers = ['Jour', 'Heure de début', 'Heure de fin', 'Salle']
        for i, header in enumerate(headers):
            cell = schedule_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(12)

        # Données du planning
        for i, sched in enumerate(student_data['schedule'], start=1):
            schedule_table.rows[i].cells[0].text = sched['day']
            schedule_table.rows[i].cells[1].text = str(sched['start_time'])
            schedule_table.rows[i].cells[2].text = str(sched['end_time'])
            schedule_table.rows[i].cells[3].text = sched['classroom']
            # Définir la taille de police à 12pt pour toutes les cellules
            for j in range(4):
                for run in schedule_table.rows[i].cells[j].paragraphs[0].runs:
                    run.font.size = Pt(12)
    else:
        para = doc.add_paragraph('Planning non défini')
        for run in para.runs:
            run.font.size = Pt(12)

    doc.add_paragraph()  # Espace
    doc.add_paragraph()  # Espace

    # Signature
    line_para = doc.add_paragraph('_' * 60)
    for run in line_para.runs:
        run.font.size = Pt(12)

    sig = doc.add_paragraph('Signature du Directeur de l\'établissement')
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sig.runs:
        run.font.size = Pt(12)

    return doc


def generate_registration_forms(group_id, output_folder='fiches_inscription', header_image='header_image.png'):
    """
    Génère les fiches d'inscription pour tous les étudiants d'un groupe.

    Args:
        group_id: ID du groupe
        output_folder: Dossier de sortie pour les fiches générées
        header_image: Chemin vers l'image d'en-tête (par défaut: 'header_image.png')
    """
    print(f"\n{'='*60}")
    print(f"Génération des fiches d'inscription pour le groupe {group_id}")
    print(f"{'='*60}\n")

    # Récupérer les données
    students_data = get_group_students_data(group_id)

    if students_data is None:
        return

    if not students_data:
        print("Aucune fiche à générer.")
        return

    # Créer le dossier de sortie
    os.makedirs(output_folder, exist_ok=True)
    print(f"\n📁 Dossier de sortie: {os.path.abspath(output_folder)}\n")

    # Vérifier si l'image d'en-tête existe
    if os.path.exists(header_image):
        print(f"✓ Image d'en-tête trouvée: {header_image}")
    else:
        print(f"⚠ Image d'en-tête non trouvée: {header_image}")
        print(f"  Les fiches seront générées avec un en-tête texte par défaut.\n")

    # Générer les fiches
    print("\nGénération des fiches Word...\n")

    for i, student_data in enumerate(students_data, 1):
        # Créer le document
        doc = create_registration_form(student_data, header_image)

        # Nom du fichier
        filename = f"Fiche_Inscription_{student_data['student_code']}_{student_data['last_name']}_{student_data['first_name']}.docx"
        filepath = os.path.join(output_folder, filename)

        # Sauvegarder
        doc.save(filepath)
        print(f"  [{i}/{len(students_data)}] ✓ {filename}")

    print(f"\n{'='*60}")
    print(f"✅ {len(students_data)} fiche(s) générée(s) avec succès!")
    print(f"📂 Emplacement: {os.path.abspath(output_folder)}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_registration_forms.py <group_id> [output_folder] [header_image]")
        print("\nExemples:")
        print("  python generate_registration_forms.py 1")
        print("  python generate_registration_forms.py 1 mes_fiches")
        print("  python generate_registration_forms.py 1 mes_fiches mon_logo.png")
        sys.exit(1)

    group_id = int(sys.argv[1])
    output_folder = sys.argv[2] if len(sys.argv) > 2 else 'fiches_inscription'
    header_image = sys.argv[3] if len(sys.argv) > 3 else 'header_image.png'

    try:
        generate_registration_forms(group_id, output_folder, header_image)
    except Exception as e:
        print(f"\n❌ Erreur: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
